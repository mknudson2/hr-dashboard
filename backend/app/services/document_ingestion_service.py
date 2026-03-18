"""
Document ingestion service for Mímir knowledge base.
Processes documents (PDF, DOCX, TXT), chunks text, embeds via OpenAI, stores in ChromaDB.
"""

import os
import uuid
import logging
from typing import Optional

from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)

# Lazy imports for optional dependencies
_chromadb = None
_openai_client = None
_text_splitter = None


def _get_chroma_collection():
    """Get or create the ChromaDB collection. Returns None if ChromaDB unavailable."""
    global _chromadb
    try:
        import chromadb
        _chromadb = chromadb

        chroma_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "chroma_data")
        os.makedirs(chroma_path, exist_ok=True)

        client = chromadb.PersistentClient(path=chroma_path)
        collection = client.get_or_create_collection(
            name="mimir_documents",
            metadata={"hnsw:space": "cosine"},
        )
        return collection
    except ImportError:
        logger.warning("chromadb not installed — document ingestion disabled")
        return None
    except Exception as e:
        logger.error(f"Failed to initialize ChromaDB: {e}")
        return None


def _get_openai_client():
    """Get OpenAI client for embeddings. Returns None if unavailable."""
    global _openai_client
    if _openai_client is not None:
        return _openai_client

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        logger.warning("OPENAI_API_KEY not set — embedding disabled, will use ChromaDB default embeddings")
        return None

    try:
        from openai import OpenAI
        _openai_client = OpenAI(api_key=api_key)
        return _openai_client
    except ImportError:
        logger.warning("openai package not installed")
        return None


def _get_text_splitter():
    """Get LangChain text splitter. Returns None if unavailable."""
    global _text_splitter
    if _text_splitter is not None:
        return _text_splitter

    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        _text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=100,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        return _text_splitter
    except ImportError:
        logger.warning("langchain-text-splitters not installed — using basic chunking")
        return None


def _basic_chunk_text(text: str, chunk_size: int = 500, overlap: int = 100) -> list[str]:
    """Fallback chunking when LangChain is not available."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start = end - overlap
    return chunks


def _extract_text_from_pdf(file_path: str) -> list[dict]:
    """Extract text from PDF, returning list of {text, page} dicts."""
    pages = []
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    pages.append({"text": text.strip(), "page": i + 1})
    except Exception as e:
        logger.error(f"Failed to extract PDF text: {e}")
    return pages


def _extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX."""
    try:
        from docx import Document
        doc = Document(file_path)
        return "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception as e:
        logger.error(f"Failed to extract DOCX text: {e}")
        return ""


def _extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception as e:
        logger.error(f"Failed to read text file: {e}")
        return ""


def _embed_texts(texts: list[str]) -> Optional[list[list[float]]]:
    """Embed texts using OpenAI. Returns None if unavailable (ChromaDB will use defaults)."""
    client = _get_openai_client()
    if client is None:
        return None

    try:
        response = client.embeddings.create(
            model="text-embedding-3-small",
            input=texts,
        )
        return [item.embedding for item in response.data]
    except Exception as e:
        logger.error(f"Failed to generate embeddings: {e}")
        return None


def ingest_document(
    file_path: str,
    original_filename: str,
    file_type: str,
    db: Session,
    uploaded_by: Optional[int] = None,
) -> Optional[dict]:
    """
    Ingest a document into the Mímir knowledge base.

    Returns dict with document_id and chunk_count, or None on failure.
    """
    from app.db.mimir_models import MimirDocument, MimirDocumentChunk

    # Extract text based on file type
    if file_type == "pdf":
        page_data = _extract_text_from_pdf(file_path)
        full_text = "\n\n".join(p["text"] for p in page_data)
    elif file_type == "docx":
        full_text = _extract_text_from_docx(file_path)
        page_data = []
    elif file_type == "txt":
        full_text = _extract_text_from_txt(file_path)
        page_data = []
    else:
        logger.error(f"Unsupported file type: {file_type}")
        return None

    if not full_text.strip():
        logger.warning(f"No text extracted from {original_filename}")
        return None

    # Chunk text
    splitter = _get_text_splitter()
    if splitter:
        chunks = splitter.split_text(full_text)
    else:
        chunks = _basic_chunk_text(full_text)

    if not chunks:
        return None

    # Get file size
    file_size = os.path.getsize(file_path) if os.path.exists(file_path) else None

    # Save document metadata to DB
    doc_record = MimirDocument(
        filename=os.path.basename(file_path),
        original_filename=original_filename,
        file_type=file_type,
        chunk_count=len(chunks),
        file_size_bytes=file_size,
        uploaded_by=uploaded_by,
    )
    db.add(doc_record)
    db.flush()  # Get the ID

    # Determine page numbers for each chunk (approximate)
    chunk_pages = []
    if page_data:
        cumulative_len = 0
        page_boundaries = []
        for p in page_data:
            cumulative_len += len(p["text"])
            page_boundaries.append((cumulative_len, p["page"]))

        for chunk in chunks:
            # Find which page this chunk is most likely from
            chunk_pos = full_text.find(chunk[:50])
            page_num = None
            for boundary_len, page in page_boundaries:
                if chunk_pos < boundary_len:
                    page_num = page
                    break
            chunk_pages.append(page_num)
    else:
        chunk_pages = [None] * len(chunks)

    # Generate embeddings
    embeddings = _embed_texts(chunks)

    # Store in ChromaDB
    collection = _get_chroma_collection()
    chunk_ids = [f"{doc_record.id}-{i}" for i in range(len(chunks))]

    if collection is not None:
        metadatas = [
            {
                "document_id": doc_record.id,
                "filename": original_filename,
                "chunk_index": i,
                "page": chunk_pages[i] or 0,
            }
            for i in range(len(chunks))
        ]

        try:
            add_kwargs = {
                "ids": chunk_ids,
                "documents": chunks,
                "metadatas": metadatas,
            }
            if embeddings is not None:
                add_kwargs["embeddings"] = embeddings

            collection.add(**add_kwargs)
        except Exception as e:
            logger.error(f"Failed to add chunks to ChromaDB: {e}")

    # Save chunk records to DB
    for i, chunk_text in enumerate(chunks):
        chunk_record = MimirDocumentChunk(
            document_id=doc_record.id,
            chunk_index=i,
            text=chunk_text,
            page=chunk_pages[i],
            embedding_id=chunk_ids[i],
        )
        db.add(chunk_record)

    db.commit()

    return {
        "document_id": doc_record.id,
        "chunk_count": len(chunks),
        "filename": original_filename,
    }


def search_documents(query: str, top_k: int = 6) -> list[dict]:
    """
    Search the Mímir knowledge base using semantic similarity.

    Returns list of {text, metadata} dicts matching the query (no similarity cutoff).
    """
    collection = _get_chroma_collection()
    if collection is None:
        return []

    try:
        # Check if collection has any documents
        if collection.count() == 0:
            return []

        # Generate query embedding if OpenAI available
        query_embedding = None
        embeddings = _embed_texts([query])
        if embeddings:
            query_embedding = embeddings[0]

        # Search ChromaDB
        if query_embedding:
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                include=["documents", "metadatas", "distances"],
            )
        else:
            # Fallback to ChromaDB's default embedding
            results = collection.query(
                query_texts=[query],
                n_results=top_k,
                include=["documents", "metadatas", "distances"],
            )

        # Process results — return all top_k chunks, no similarity cutoff
        chunks = []
        if results and results["documents"] and results["documents"][0]:
            for i, doc_text in enumerate(results["documents"][0]):
                distance = results["distances"][0][i] if results["distances"] else 0
                similarity = 1 - distance
                metadata = results["metadatas"][0][i] if results["metadatas"] else {}
                chunks.append({
                    "text": doc_text,
                    "metadata": metadata,
                    "similarity": similarity,
                })

        return chunks

    except Exception as e:
        logger.error(f"Failed to search ChromaDB: {e}")
        return []
