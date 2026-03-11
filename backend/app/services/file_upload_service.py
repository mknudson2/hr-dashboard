"""
File Upload and Processing Service
Handles secure file uploads, validation, parsing, and data import
"""

import os
import uuid
import pandas as pd
from datetime import datetime
from typing import Dict, List, Tuple, Any, Optional
from pathlib import Path
import aiofiles
from sqlalchemy.orm import Session
from fastapi import UploadFile, HTTPException
import logging

# PDF processing
import PyPDF2
import pdfplumber

# Word document processing
from docx import Document as DocxDocument

# Database models
from app.db import models

# Configure logging FIRST
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Try to import python-magic, but make it optional
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    logger.warning("python-magic not available. MIME type detection will use extension-based validation only.")

# ============================================================================
# SECURITY CONFIGURATION
# ============================================================================

# File size limits (in bytes)
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB

# Allowed file types and their MIME types
ALLOWED_FILE_TYPES = {
    'csv': ['text/csv', 'application/csv', 'text/plain'],
    'xlsx': ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'],
    'xls': ['application/vnd.ms-excel'],
    'docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
    'pdf': ['application/pdf'],
}

# Upload directory (outside web root for security)
UPLOAD_DIR = Path(__file__).parent.parent / "data" / "uploads"
QUARANTINE_DIR = UPLOAD_DIR / "quarantine"

# Ensure directories exist
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
QUARANTINE_DIR.mkdir(parents=True, exist_ok=True)


# ============================================================================
# FILE VALIDATION
# ============================================================================

class FileValidator:
    """Validates uploaded files for security and compliance"""

    @staticmethod
    def validate_file_size(file_size: int) -> Tuple[bool, Optional[str]]:
        """Validate file size is within limits"""
        if file_size > MAX_FILE_SIZE:
            return False, f"File size ({file_size} bytes) exceeds maximum allowed size ({MAX_FILE_SIZE} bytes)"
        if file_size == 0:
            return False, "File is empty"
        return True, None

    @staticmethod
    def validate_file_extension(filename: str) -> Tuple[bool, Optional[str], Optional[str]]:
        """Validate file has an allowed extension"""
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else None
        if not ext or ext not in ALLOWED_FILE_TYPES:
            allowed = ', '.join(ALLOWED_FILE_TYPES.keys())
            return False, f"File type '.{ext}' not allowed. Allowed types: {allowed}", None
        return True, None, ext

    @staticmethod
    async def validate_mime_type(file_path: str, expected_type: str) -> Tuple[bool, Optional[str]]:
        """Validate file MIME type using magic bytes (prevents extension spoofing)"""
        if not MAGIC_AVAILABLE:
            # If magic is not available, skip MIME validation (rely on extension only)
            logger.warning(f"Skipping MIME validation for {expected_type} (python-magic not available)")
            return True, None

        try:
            mime = magic.Magic(mime=True)
            detected_mime = mime.from_file(file_path)

            allowed_mimes = ALLOWED_FILE_TYPES.get(expected_type, [])
            if detected_mime not in allowed_mimes:
                return False, f"File MIME type '{detected_mime}' does not match expected type for .{expected_type} ({', '.join(allowed_mimes)})"
            return True, None
        except Exception as e:
            logger.error(f"Error validating MIME type: {e}")
            # Don't fail upload if MIME validation fails, just log warning
            logger.warning(f"MIME validation failed, allowing file based on extension only")
            return True, None

    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """Sanitize filename to prevent directory traversal attacks"""
        # Remove path components
        filename = os.path.basename(filename)
        # Remove dangerous characters
        filename = "".join(c for c in filename if c.isalnum() or c in ('_', '-', '.'))
        return filename

    @staticmethod
    async def scan_for_malware(file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Hook for malware scanning (ClamAV integration)
        Currently returns True (no scan), but can be extended
        """
        # Malware scanning can be enabled via ClamAV integration
        return True, None


# ============================================================================
# FILE STORAGE
# ============================================================================

class FileStorage:
    """Handles secure file storage"""

    @staticmethod
    async def save_uploaded_file(
        upload_file: UploadFile,
        uploaded_by: str
    ) -> Tuple[str, str, int]:
        """
        Save uploaded file securely with UUID naming
        Returns: (file_path, secure_filename, file_size)
        """
        # Generate UUID-based filename
        ext = upload_file.filename.rsplit('.', 1)[-1].lower() if '.' in upload_file.filename else ''
        secure_filename = f"{uuid.uuid4()}.{ext}"

        # Create date-based subdirectory for organization
        today = datetime.now()
        subdir = UPLOAD_DIR / str(today.year) / f"{today.month:02d}"
        subdir.mkdir(parents=True, exist_ok=True)

        file_path = subdir / secure_filename

        # Write file securely
        file_size = 0
        async with aiofiles.open(file_path, 'wb') as f:
            while content := await upload_file.read(8192):  # Read in chunks
                file_size += len(content)
                await f.write(content)

        # Set secure file permissions (read/write for owner only)
        os.chmod(file_path, 0o600)

        return str(file_path), secure_filename, file_size

    @staticmethod
    async def quarantine_file(file_path: str, reason: str):
        """Move suspicious file to quarantine directory"""
        filename = os.path.basename(file_path)
        quarantine_path = QUARANTINE_DIR / f"{datetime.now().isoformat()}_{filename}"

        try:
            os.rename(file_path, quarantine_path)
            logger.warning(f"File quarantined: {filename} - Reason: {reason}")
        except Exception as e:
            logger.error(f"Error quarantining file: {e}")

    @staticmethod
    async def delete_file(file_path: str):
        """Securely delete a file"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"File deleted: {file_path}")
        except Exception as e:
            logger.error(f"Error deleting file: {e}")


# ============================================================================
# FILE PARSERS
# ============================================================================

class CSVParser:
    """Parse and validate CSV files"""

    @staticmethod
    async def parse(file_path: str) -> Tuple[pd.DataFrame, List[Dict]]:
        """
        Parse CSV file and return dataframe + logs
        Returns: (dataframe, logs)
        """
        logs = []
        try:
            # Read CSV with pandas
            df = pd.read_csv(file_path)

            logs.append({
                'level': 'info',
                'message': f'Successfully parsed CSV file with {len(df)} rows and {len(df.columns)} columns',
                'details': {'rows': len(df), 'columns': list(df.columns)}
            })

            return df, logs

        except pd.errors.EmptyDataError:
            logs.append({
                'level': 'error',
                'message': 'CSV file is empty',
                'details': {}
            })
            raise HTTPException(status_code=400, detail="CSV file is empty")

        except Exception as e:
            logs.append({
                'level': 'error',
                'message': f'Error parsing CSV: {str(e)}',
                'details': {'error': str(e)}
            })
            raise HTTPException(status_code=400, detail=f"Error parsing CSV: {str(e)}")


class ExcelParser:
    """Parse and validate Excel files"""

    @staticmethod
    async def parse(file_path: str, sheet_name: Optional[str] = None) -> Tuple[pd.DataFrame, List[Dict]]:
        """
        Parse Excel file and return dataframe + logs
        Returns: (dataframe, logs)
        """
        logs = []
        try:
            # Read Excel with pandas — try openpyxl first, fall back to calamine
            # for Strict Open XML files that openpyxl can't handle
            try:
                if sheet_name:
                    df = pd.read_excel(file_path, sheet_name=sheet_name, engine='openpyxl')
                else:
                    df = pd.read_excel(file_path, engine='openpyxl')
            except Exception:
                if sheet_name:
                    df = pd.read_excel(file_path, sheet_name=sheet_name, engine='calamine')
                else:
                    df = pd.read_excel(file_path, engine='calamine')

            logs.append({
                'level': 'info',
                'message': f'Successfully parsed Excel file with {len(df)} rows and {len(df.columns)} columns',
                'details': {'rows': len(df), 'columns': list(df.columns), 'sheet': sheet_name or 'default'}
            })

            return df, logs

        except Exception as e:
            logs.append({
                'level': 'error',
                'message': f'Error parsing Excel: {str(e)}',
                'details': {'error': str(e)}
            })
            raise HTTPException(status_code=400, detail=f"Error parsing Excel: {str(e)}")


class PDFParser:
    """Parse and extract text from PDF files"""

    @staticmethod
    async def parse(file_path: str) -> Tuple[str, List[Dict]]:
        """
        Extract text from PDF file
        Returns: (extracted_text, logs)
        """
        logs = []
        extracted_text = ""

        try:
            # Try pdfplumber first (better for tables)
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                logs.append({
                    'level': 'info',
                    'message': f'PDF has {page_count} pages',
                    'details': {'pages': page_count}
                })

                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        extracted_text += f"\n--- Page {i+1} ---\n{text}\n"

            logs.append({
                'level': 'info',
                'message': f'Successfully extracted text from PDF ({len(extracted_text)} characters)',
                'details': {'text_length': len(extracted_text)}
            })

            return extracted_text, logs

        except Exception as e:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    page_count = len(pdf_reader.pages)

                    for i, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        if text:
                            extracted_text += f"\n--- Page {i+1} ---\n{text}\n"

                logs.append({
                    'level': 'warning',
                    'message': f'Used fallback PDF parser. Extracted {len(extracted_text)} characters',
                    'details': {'text_length': len(extracted_text), 'pages': page_count}
                })

                return extracted_text, logs

            except Exception as e2:
                logs.append({
                    'level': 'error',
                    'message': f'Error parsing PDF: {str(e2)}',
                    'details': {'error': str(e2)}
                })
                raise HTTPException(status_code=400, detail=f"Error parsing PDF: {str(e2)}")


class DocxParser:
    """Parse and extract text from Word documents"""

    @staticmethod
    async def parse(file_path: str) -> Tuple[str, List[Dict]]:
        """
        Extract text from DOCX file
        Returns: (extracted_text, logs)
        """
        logs = []
        extracted_text = ""

        try:
            doc = DocxDocument(file_path)

            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            extracted_text = "\n".join(paragraphs)

            # Extract table data
            table_count = len(doc.tables)
            if table_count > 0:
                extracted_text += "\n\n--- Tables ---\n"
                for i, table in enumerate(doc.tables):
                    extracted_text += f"\nTable {i+1}:\n"
                    for row in table.rows:
                        row_text = " | ".join(cell.text for cell in row.cells)
                        extracted_text += row_text + "\n"

            logs.append({
                'level': 'info',
                'message': f'Successfully extracted text from DOCX ({len(paragraphs)} paragraphs, {table_count} tables)',
                'details': {
                    'paragraphs': len(paragraphs),
                    'tables': table_count,
                    'text_length': len(extracted_text)
                }
            })

            return extracted_text, logs

        except Exception as e:
            logs.append({
                'level': 'error',
                'message': f'Error parsing DOCX: {str(e)}',
                'details': {'error': str(e)}
            })
            raise HTTPException(status_code=400, detail=f"Error parsing DOCX: {str(e)}")


# ============================================================================
# MAIN FILE UPLOAD SERVICE
# ============================================================================

class FileUploadService:
    """Main service for handling file uploads"""

    @staticmethod
    async def upload_and_validate(
        upload_file: UploadFile,
        uploaded_by: str,
        db: Session,
        file_category: Optional[str] = None
    ) -> models.FileUpload:
        """
        Upload file, validate it, and create database record
        Returns: FileUpload model instance

        Args:
            upload_file: The uploaded file
            uploaded_by: Username of uploader
            db: Database session
            file_category: Optional category hint (employment_list, ot_earnings, etc.)
        """
        logs = []

        try:
            # Step 1: Validate filename and extension
            original_filename = FileValidator.sanitize_filename(upload_file.filename)
            is_valid, error, file_type = FileValidator.validate_file_extension(original_filename)
            if not is_valid:
                raise HTTPException(status_code=400, detail=error)

            # Step 2: Save file temporarily to check size and MIME type
            file_path, secure_filename, file_size = await FileStorage.save_uploaded_file(
                upload_file, uploaded_by
            )

            # Step 3: Validate file size
            is_valid, error = FileValidator.validate_file_size(file_size)
            if not is_valid:
                await FileStorage.delete_file(file_path)
                raise HTTPException(status_code=400, detail=error)

            # Step 4: Validate MIME type (prevents extension spoofing)
            is_valid, error = await FileValidator.validate_mime_type(file_path, file_type)
            if not is_valid:
                await FileStorage.quarantine_file(file_path, error)
                raise HTTPException(status_code=400, detail=error)

            # Step 5: Scan for malware (hook for future ClamAV integration)
            is_safe, error = await FileValidator.scan_for_malware(file_path)
            if not is_safe:
                await FileStorage.quarantine_file(file_path, error)
                raise HTTPException(status_code=400, detail=f"Security scan failed: {error}")

            # Step 6: Detect MIME type for storage
            if MAGIC_AVAILABLE:
                mime = magic.Magic(mime=True)
                detected_mime = mime.from_file(file_path)
            else:
                # Fallback to extension-based MIME type
                detected_mime = ALLOWED_FILE_TYPES.get(file_type, ['application/octet-stream'])[0]

            # Step 7: Create database record
            file_upload = models.FileUpload(
                file_name=secure_filename,
                original_filename=original_filename,
                file_type=file_type,
                file_size=file_size,
                file_path=file_path,
                mime_type=detected_mime,
                upload_source='manual',
                uploaded_by=uploaded_by,
                status='pending',
                file_category=file_category  # Store category hint
            )

            db.add(file_upload)
            db.commit()
            db.refresh(file_upload)

            logger.info(f"File uploaded successfully: {original_filename} -> {secure_filename}")

            return file_upload

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Unexpected error during file upload: {e}")
            raise HTTPException(status_code=500, detail=f"File upload failed: {str(e)}")

    @staticmethod
    async def parse_and_process(
        file_upload_id: int,
        db: Session,
        force_category: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Parse an uploaded file and process the data

        Args:
            file_upload_id: ID of the uploaded file
            db: Database session
            force_category: Force a specific category instead of auto-detect

        Returns:
            Dictionary with parsing results and statistics
        """
        from app.services.file_parsers import AutoDetectParser, get_parser, FileCategory

        # Get file upload record
        file_upload = db.query(models.FileUpload).filter(
            models.FileUpload.id == file_upload_id
        ).first()

        if not file_upload:
            raise HTTPException(status_code=404, detail="File upload not found")

        try:
            # Update status
            file_upload.status = 'processing'
            file_upload.processing_started_at = datetime.now()
            db.commit()

            # Parse file
            if force_category:
                # Use specific parser
                category = FileCategory(force_category)
                parser_class = get_parser(category)
                if not parser_class:
                    raise HTTPException(status_code=400, detail=f"No parser available for category: {force_category}")
                df, config, logs = await parser_class.parse(file_upload.file_path)
            else:
                # Auto-detect
                df, config, logs = await AutoDetectParser.parse(file_upload.file_path)

            # Update file upload with category
            file_upload.file_category = config.category.value
            file_upload.detected_columns = df.columns.tolist()
            file_upload.row_count = len(df)

            # Store processing logs
            file_upload.processing_logs = logs

            # Update status
            file_upload.status = 'parsed'
            file_upload.processing_completed_at = datetime.now()

            db.commit()
            db.refresh(file_upload)

            return {
                'file_upload_id': file_upload.id,
                'category': config.category.value,
                'category_name': config.name,
                'rows': len(df),
                'columns': len(df.columns),
                'logs': logs,
                'data_preview': df.head(10).to_dict('records') if len(df) > 0 else []
            }

        except HTTPException:
            file_upload.status = 'failed'
            file_upload.processing_completed_at = datetime.now()
            db.commit()
            raise
        except Exception as e:
            file_upload.status = 'failed'
            file_upload.processing_completed_at = datetime.now()
            file_upload.processing_logs = [{
                'level': 'error',
                'message': str(e),
                'details': {}
            }]
            db.commit()
            logger.error(f"Error processing file {file_upload_id}: {e}")
            raise HTTPException(status_code=500, detail=f"File processing failed: {str(e)}")
